from pathlib import Path

import pytest

from fileconv.docs import (
    _docx_to_pdf_pandoc,
    _find_libreoffice,
    _pandoc_typst_available,
    _word_available,
    docx_to_pdf,
    pdf_to_docx,
)


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """One-page PDF generated with PyMuPDF (already installed via pdf2docx)."""
    import fitz

    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from fileconv test suite.")
    doc.save(str(path))
    doc.close()
    return path


def test_pdf_to_docx(sample_pdf: Path, tmp_path: Path):
    dest = pdf_to_docx(sample_pdf, tmp_path / "out.docx")
    assert dest.exists() and dest.stat().st_size > 0


@pytest.mark.skipif(
    not (_word_available() or _find_libreoffice() or _pandoc_typst_available()),
    reason="no DOCX -> PDF engine available",
)
def test_docx_to_pdf_round_trip(sample_pdf: Path, tmp_path: Path):
    docx = pdf_to_docx(sample_pdf, tmp_path / "mid.docx")
    pdf = docx_to_pdf(docx, tmp_path / "back.pdf")
    assert pdf.exists() and pdf.stat().st_size > 0


@pytest.mark.skipif(not _pandoc_typst_available(), reason="pandoc/typst not installed")
def test_pandoc_engine_with_formatting_and_image(tmp_path: Path):
    """Exercise the pure-Python engine directly on a DOCX containing headings,
    bold text, a list, a table, and an embedded image."""
    import fitz
    from PIL import Image

    img_path = tmp_path / "pic.png"
    Image.new("RGB", (40, 30), (60, 90, 200)).save(img_path)

    from docx import Document  # python-docx, installed via pdf2docx

    doc = Document()
    doc.add_heading("Test Document", level=1)
    para = doc.add_paragraph("Plain text with ")
    para.add_run("bold words").bold = True
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Second bullet", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(1, 1).text = "B"
    doc.add_picture(str(img_path))
    docx_path = tmp_path / "rich.docx"
    doc.save(str(docx_path))

    pdf = _docx_to_pdf_pandoc(docx_path, tmp_path / "rich.pdf")
    assert pdf.exists() and pdf.stat().st_size > 0
    with fitz.open(str(pdf)) as result:
        text = result[0].get_text()
    assert "Test Document" in text
    assert "bold words" in text
